import os
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches
from src.config import Config
from src.utils import setup_logging, generate_variable_symbol, calculate_duration_days

logger = setup_logging()

class DocumentEngine:
    def __init__(self):
        os.makedirs('output', exist_ok=True)
    
    def generate_documents(self, extracted_data, request_id):
        """Generate consent document and payment instructions"""
        documents = {}
        
        # Generate Variable Symbol
        vs = generate_variable_symbol(request_id)
        
        # Calculate duration and fee
        duration_days = self._calculate_duration(extracted_data.get('duration', {}))
        area_sqm = extracted_data.get('area_square_meters') or extracted_data.get('area_sqm') or extracted_data.get('area', 0)
        fee = area_sqm * duration_days * Config.DEFAULT_RATE_PER_SQM_DAY if area_sqm else 0
        
        # Add calculated values to data
        extracted_data['variable_symbol'] = vs
        extracted_data['duration_days'] = duration_days
        extracted_data['fee_czk'] = int(fee)
        
        # Generate consent document
        consent_path = self._generate_consent(extracted_data, request_id)
        documents['consent'] = consent_path
        
        # Generate payment instructions
        payment_path = self._generate_payment_instructions(extracted_data, request_id)
        documents['payment'] = payment_path
        
        return documents
    
    def _generate_consent(self, data, request_id):
        """Generate consent/permit document"""
        doc = Document()
        
        # Header
        header = doc.add_heading('SOUHLAS K ZVLÁŠTNÍMU UŽÍVÁNÍ VEŘEJNÉHO PROSTRANSTVÍ', 0)
        header.alignment = 1  # Center alignment
        
        doc.add_paragraph()
        
        # Document content
        doc.add_paragraph(f"Číslo žádosti: {request_id}")
        doc.add_paragraph(f"Datum vystavení: {datetime.now().strftime('%d.%m.%Y')}")
        
        doc.add_paragraph()
        
        # Applicant details - handle both field name formats
        doc.add_heading('Údaje žadatele:', level=2)
        applicant_name = data.get('applicant_name') or data.get('Applicant name', 'N/A')
        company_id = data.get('company_id') or data.get('Company ID (IČO)', '')
        contact = data.get('contact_details') or data.get('Contact details', '')
        
        doc.add_paragraph(f"Jméno/Název: {applicant_name}")
        if company_id:
            doc.add_paragraph(f"IČO: {company_id}")
        if contact:
            doc.add_paragraph(f"Kontakt: {contact}")
        
        doc.add_paragraph()
        
        # Usage details
        doc.add_heading('Údaje o užívání:', level=2)
        purpose = data.get('purpose_of_use') or data.get('purpose') or data.get('Purpose of use', 'N/A')
        location = data.get('specific_location') or data.get('location') or data.get('Location', 'N/A')
        duration = data.get('duration') or data.get('Duration (dates)', 'N/A')
        
        doc.add_paragraph(f"Účel užívání: {purpose}")
        doc.add_paragraph(f"Místo: {location}")
        doc.add_paragraph(f"Doba užívání: {duration}")
        
        # Fee calculation
        area = data.get('area_square_meters') or data.get('area_sqm') or data.get('area', 0)
        if area:
            doc.add_paragraph(f"Výměra: {area} m²")
            doc.add_paragraph(f"Poplatek: {data.get('fee_czk', 0)} Kč")
        
        doc.add_paragraph()
        
        # Standard conditions
        doc.add_heading('Podmínky:', level=2)
        conditions = [
            "Žadatel je povinen dodržovat všechny platné právní předpisy.",
            "Užívání je povoleno pouze v uvedeném rozsahu a době.",
            "Žadatel odpovídá za případné škody způsobené užíváním.",
            "Poplatek je splatný do 30 dnů od vystavení tohoto souhlasu."
        ]
        
        for condition in conditions:
            doc.add_paragraph(f"• {condition}")
        
        # Save document
        output_path = f"output/consent_{request_id}.docx"
        doc.save(output_path)
        logger.info(f"Consent document generated: {output_path}")
        
        return output_path
    
    def _generate_payment_instructions(self, data, request_id):
        """Generate payment instructions document"""
        doc = Document()
        
        doc.add_heading('PLATEBNÍ INSTRUKCE', 0).alignment = 1
        doc.add_paragraph()
        
        # Calculate fee
        area = data.get('area_square_meters') or data.get('area_sqm') or data.get('area', 0)
        duration = data.get('duration', {})
        fee = data.get('fee_czk', 0)  # Use pre-calculated fee
        
        # Variable symbol (using request_id)
        vs = request_id.replace('-', '')[:10]  # Use first 10 chars of request_id
        
        # Payment details
        doc.add_paragraph(f"Číslo žádosti: {request_id}")
        doc.add_paragraph(f"Částka k úhradě: {data.get('fee_czk', 0)} Kč")
        doc.add_paragraph(f"Variabilní symbol: {data.get('variable_symbol', vs)}")
        doc.add_paragraph("Číslo účtu: 123456789/0100")  # Configurable
        doc.add_paragraph(f"Splatnost: {(datetime.now() + timedelta(days=30)).strftime('%d.%m.%Y')}")
        
        doc.add_paragraph()
        doc.add_paragraph("Prosím uhraďte poplatek ve stanovené lhůtě.")
        
        # Save document
        output_path = f"output/payment_{request_id}.docx"
        doc.save(output_path)
        logger.info(f"Payment instructions generated: {output_path}")
        
        return output_path
    
    def _calculate_duration(self, duration_data):
        """Calculate duration in days from duration data"""
        try:
            if isinstance(duration_data, str):
                # Parse duration string like "24.12.2025 - 31.12.2025"
                if ' - ' in duration_data:
                    start_str, end_str = duration_data.split(' - ')
                    return calculate_duration_days(start_str.strip(), end_str.strip())
            elif isinstance(duration_data, dict):
                # Handle both start/end and start_date/end_date formats
                start_date = duration_data.get('start_date') or duration_data.get('start')
                end_date = duration_data.get('end_date') or duration_data.get('end')
                if start_date and end_date:
                    return calculate_duration_days(start_date, end_date)
            return 7  # Default fallback
        except Exception as e:
            logger.warning(f"Duration calculation failed: {e}")
            return 7
    
    def _calculate_fee(self, area, duration):
        """Calculate fee based on area and duration"""
        try:
            if isinstance(duration, str) and ' - ' in duration:
                duration_days = self._calculate_duration(duration)
            elif isinstance(duration, (int, float)):
                duration_days = duration
            else:
                duration_days = 7  # Default
            
            return int(area * duration_days * Config.DEFAULT_RATE_PER_SQM_DAY)
        except:
            return 0